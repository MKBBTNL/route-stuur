#!/usr/bin/env python3
"""
scan_manuals.py - extract text + tables from Word ops manuals (locally) and
cross-reference your uncovered route names against them, so you learn WHICH
manual/tour mentions each route instead of getting a pile of txt.

    # just extract everything to searchable text:
    python scan_manuals.py --dir "C:\\path\\to\\ops manuals" --out manuals_index.csv --dump-txt

    # cross-reference uncovered routes -> which manual mentions them:
    python scan_manuals.py --dir "C:\\path\\to\\ops manuals" --routes coverage_out/unassigned_EN.csv --hits manual_hits.csv

Needs:  python -m pip install python-docx
Reads .docx (paragraphs AND table cells). Legacy .doc is skipped with a note
(re-save those as .docx, or convert in bulk, then re-run).
"""
import argparse, csv, glob, os, re, sys

try:
    from docx import Document
except ImportError:
    sys.exit("Missing dependency. Run:  python -m pip install python-docx")

def norm(s):
    s = str(s).lower()
    s = s.replace("–","-").replace("—","-").replace("|"," ")   # unify dashes/pipes
    return re.sub(r"\s+"," ", s).strip()

def stem(name):
    s = str(name)
    s = re.sub(r"\[[^\]]*\]"," ",s)
    s = re.sub(r"\b20\d{2}\b"," ",s)
    s = re.sub(r"\b(ENG|EN|ENGLISH|NL|DUTCH|NEDERLANDS?|DE|DEU|GER|DEUTSCH|FR|FRENCH)\b"," ",s,flags=re.I)
    s = re.sub(r"\b\d+\s*km\b"," ",s,flags=re.I)
    s = re.sub(r"\b\d+\b"," ",s)
    return norm(s)

def extract_docx(path):
    """Return (full_text, n_paragraphs, n_tables). Includes table cell text."""
    doc = Document(path)
    parts, ntab = [], 0
    for p in doc.paragraphs:
        if p.text.strip(): parts.append(p.text)
    for t in doc.tables:
        ntab += 1
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells): parts.append(" | ".join(cells))
    return "\n".join(parts), len(doc.paragraphs), ntab

def load_route_names(path):
    names=[]
    with open(path, encoding="utf-8-sig", newline="") as f:
        r=csv.reader(f)
        header=next(r, [])
        # find a 'name' column, else use first column
        idx=0
        for i,h in enumerate(header):
            if h.strip().lower()=="name": idx=i; break
        else:
            # header had no 'name' -> maybe it *was* a name; keep it
            if header and header[0].strip(): names.append(header[0].strip())
        for row in r:
            if len(row)>idx and row[idx].strip(): names.append(row[idx].strip())
    # de-dupe, keep order
    seen=set(); out=[]
    for n in names:
        if n not in seen: seen.add(n); out.append(n)
    return out

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument("--dir", required=True, help="folder of Word manuals (searched recursively)")
    ap.add_argument("--routes", default=None, help="CSV with a 'name' column (e.g. coverage_out/unassigned_EN.csv)")
    ap.add_argument("--out", default="manuals_index.csv")
    ap.add_argument("--hits", default="manual_hits.csv")
    ap.add_argument("--dump-txt", action="store_true", help="also write a .txt next to each manual")
    a=ap.parse_args()
    if not os.path.isdir(a.dir): sys.exit(f"folder not found: {a.dir!r}")

    docx = glob.glob(os.path.join(a.dir,"**","*.docx"), recursive=True)
    legacy = glob.glob(os.path.join(a.dir,"**","*.doc"), recursive=True)
    print(f"found {len(docx)} .docx  ({len(legacy)} legacy .doc will be skipped)")

    corpus=[]   # (relpath, tourguess, text, ntext_norm)
    index=[]
    for fp in sorted(docx):
        try:
            text, npar, ntab = extract_docx(fp)
        except Exception as e:
            index.append({"file":os.path.basename(fp),"path":fp,"paragraphs":"","tables":"","chars":"","error":str(e)[:80]})
            continue
        rel=os.path.relpath(fp, a.dir)
        tour=os.path.splitext(os.path.basename(fp))[0]
        corpus.append((rel, tour, text, norm(text)))
        index.append({"file":os.path.basename(fp),"path":rel,"paragraphs":npar,"tables":ntab,"chars":len(text),"error":""})
        if a.dump_txt:
            open(os.path.splitext(fp)[0]+".txt","w",encoding="utf-8").write(text)

    with open(a.out,"w",newline="",encoding="utf-8") as f:
        w=csv.DictWriter(f, fieldnames=["file","path","paragraphs","tables","chars","error"]); w.writeheader()
        for r in index: w.writerow(r)
    print(f"wrote {a.out}  ({len(index)} manuals indexed)")

    if not a.routes:
        print("no --routes given; extraction only. Re-run with --routes to cross-reference.")
        return
    names=load_route_names(a.routes)
    print(f"cross-referencing {len(names)} route names…")

    hits=[]; found=set()
    for nm in names:
        key=stem(nm)
        if len(key)<6:   # too short to match safely
            continue
        for rel,tour,text,tnorm in corpus:
            pos=tnorm.find(key)
            if pos!=-1:
                snip=tnorm[max(0,pos-60):pos+len(key)+60]
                hits.append({"route_name":nm,"stem":key,"manual":rel,"tour_guess":tour,"snippet":"…"+snip+"…"})
                found.add(nm)
    with open(a.hits,"w",newline="",encoding="utf-8") as f:
        w=csv.DictWriter(f, fieldnames=["route_name","stem","manual","tour_guess","snippet"]); w.writeheader()
        for r in hits: w.writerow(r)
    print(f"wrote {a.hits}  ({len(hits)} mentions; {len(found)}/{len(names)} routes found in a manual)")
    miss=[n for n in names if n not in found]
    if miss:
        print(f"{len(miss)} routes had NO mention in any manual (sample):")
        for n in miss[:8]: print("  -",n)

if __name__=="__main__": main()
